from fastapi import FastAPI, UploadFile, File, Header
from pydantic import BaseModel
import os
import pandas as pd
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.llms import Ollama
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders import PyPDFLoader, TextLoader, CSVLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import tempfile
from collections import defaultdict

# Import Chroma from langchain_chroma
from langchain_chroma import Chroma

# Initialize FastAPI app
app = FastAPI()

# In-memory storage for user sessions and chat history
user_sessions = defaultdict(lambda: defaultdict(list))



# Function to extract text from a PDF file
def extract_text_from_pdf(pdf_path):
    try:
        loader = PyPDFLoader(pdf_path)
        pages = loader.load_and_split()
        text = "\n".join(page.page_content for page in pages)
        return text
    except Exception:
        return ""

# Function to extract text from a DOCX file
def extract_text_from_docx(docx_path):
    try:
        import docx
        doc = docx.Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception:
        return ""

# Function to extract text from a TXT file
def extract_text_from_txt(txt_path):
    try:
        loader = TextLoader(txt_path)
        documents = loader.load()
        text = "\n".join(doc.page_content for doc in documents)
        return text
    except Exception:
        return ""

# Function to extract text from a CSV file
def extract_text_from_csv(csv_path):
    try:
        loader = CSVLoader(file_path=csv_path)
        data = loader.load()
        text = "\n".join([str(record) for record in data])
        return text
    except Exception:
        return ""

# Function to extract text from an XLSX file
def extract_text_from_xlsx(xlsx_path):
    try:
        df = pd.read_excel(xlsx_path)
        text = df.to_string(index=False)
        return text
    except Exception:
        return ""

# Function to get file extension
def get_file_extension(file_path):
    return os.path.splitext(file_path)[1].lower()

# Function to extract text based on file type
def extract_text(file_path):
    ext = get_file_extension(file_path)
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    elif ext == '.csv':
        return extract_text_from_csv(file_path)
    elif ext == '.xlsx':
        return extract_text_from_xlsx(file_path)
    else:
        return ""
    
# Initialize LLM
llm = Ollama(model="cniongolo/biomistral")

# Define the prompt template
prompt_template = """Use the following pieces of information to answer the user's question.
If you don't know the answer, just say that you don't know, don't try to make up an answer.

Context: {context}
Question: {question}
Chat History: {chat_history}

Only return the helpful answer below and nothing else and no questions to be asked again.
Helpful answer:
"""

# Initialize the Hugging Face Embeddings
embeddings = HuggingFaceEmbeddings(model_name="NeuML/pubmedbert-base-embeddings")

# Directory for Chroma DB (company data)
company_data_directory = "chroma_db"

# Initialize vector store for company data
company_vector_store = Chroma(persist_directory=company_data_directory, embedding_function=embeddings)

# Setup prompt
prompt = PromptTemplate(template=prompt_template, input_variables=['context', 'question', 'chat_history'])

# API to receive question and return LLM response
class QuestionRequest(BaseModel):
    query: str

@app.post("/chat/")
async def ask_question(request: QuestionRequest, x_user_id: str = Header(...), x_session_id: str = Header(...)):
    question = request.query
    user_id = x_user_id
    session_id = x_session_id

    try:
        # Retrieve relevant documents from company data vector store
        retriever = company_vector_store.as_retriever(search_kwargs={"k": 3})
        docs = retriever.invoke(question)  # Updated method

        if not docs:
            return {"response": "No relevant documents found."}

        context = "\n".join([doc.page_content for doc in docs])

        # Get chat history for the current session
        chat_history = "\n".join(user_sessions[user_id][session_id])

        # Generate the formatted prompt
        formatted_prompt = prompt.format(context=context, question=question, chat_history=chat_history)

        # Get response from the LLM
        response = llm.invoke(formatted_prompt, max_token=2048)  # Updated method

        # Store the current question and response in chat history
        user_sessions[user_id][session_id].append(f"Q: {question}\nA: {response.strip()}")

        # Return the LLM response only
        return {"response": response.strip()}

    except Exception as e:
        return {"response": f"Error: {str(e)}"}

# API to upload files and ingest them into the company database
@app.post("/ingest_file/")
async def ingest_file(file: UploadFile = File(...)):
    try:
        # Save the uploaded file to a temporary location
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(await file.read())
            temp_file_path = temp_file.name

        # Extract text based on file type
        text = extract_text(temp_file_path)

        if not text:
            return {"status": "Failed to extract text"}

        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        documents = [Document(page_content=text)]
        texts = text_splitter.split_documents(documents)

        # Ingest the content into the company vector store
        texts_content = [doc.page_content for doc in texts]
        company_vector_store.add_texts(texts_content)

        # Clean up the temporary file
        os.remove(temp_file_path)

        return {"status": "File ingested successfully"}

    except Exception as e:
        return {"status": f"Error: {str(e)}"}
    
@app.get("/")
async def root():
    return {"message": "Hello World"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
